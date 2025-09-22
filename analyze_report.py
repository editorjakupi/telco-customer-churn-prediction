#!/usr/bin/env python3
"""
Analyze the report content and compare with actual project results
"""

from docx import Document
import json
import re

def get_actual_results():
    """Get actual results from model_info.json"""
    try:
        with open('model_info.json', 'r') as f:
            model_info = json.load(f)
        return model_info
    except Exception as e:
        print(f"Error loading model info: {e}")
        return None

def analyze_report_content():
    """Analyze the report content"""
    try:
        # Load the document
        doc = Document('individual_report.docx')
        
        # Get actual results
        actual_results = get_actual_results()
        if not actual_results:
            print("Could not load actual results")
            return
        
        print("=== ACTUAL PROJECT RESULTS ===")
        print(f"Best Model: {actual_results.get('best_model', 'N/A')}")
        print(f"Validation Accuracy: {actual_results.get('best_accuracy', 0):.3f}")
        print(f"Test Accuracy: {actual_results.get('test_accuracy', 0):.3f}")
        print(f"Total Features: {actual_results.get('total_features', 0)}")
        print(f"Numerical Columns: {actual_results.get('numerical_columns', [])}")
        print(f"Categorical Columns: {len(actual_results.get('categorical_columns', []))} columns")
        print(f"Engineered Features: {actual_results.get('engineered_features', [])}")
        
        print("\n=== REPORT CONTENT ANALYSIS ===")
        
        # Check for accuracy values in report
        accuracy_found = False
        model_found = False
        features_found = False
        preprocessing_found = False
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Check for accuracy values
            if re.search(r'74\.\d+%|0\.74\d+', text):
                print(f"Found accuracy in paragraph {i}: {text[:100]}...")
                accuracy_found = True
            
            # Check for model mentions
            if 'Random Forest' in text and 'model' in text.lower():
                print(f"Found model mention in paragraph {i}: {text[:100]}...")
                model_found = True
            
            # Check for feature count
            if re.search(r'\d+\s+features?', text):
                print(f"Found feature count in paragraph {i}: {text[:100]}...")
                features_found = True
            
            # Check for preprocessing methods
            if any(method in text for method in ['ColumnTransformer', 'StandardScaler', 'OneHotEncoder']):
                print(f"Found preprocessing method in paragraph {i}: {text[:100]}...")
                preprocessing_found = True
        
        print("\n=== SYNC ANALYSIS ===")
        print(f"✅ Accuracy values found: {accuracy_found}")
        print(f"✅ Model mentioned: {model_found}")
        print(f"✅ Feature count mentioned: {features_found}")
        print(f"✅ Preprocessing methods mentioned: {preprocessing_found}")
        
        # Check for specific issues
        issues = []
        
        # Check for LabelEncoder (should not be there)
        for paragraph in doc.paragraphs:
            if 'LabelEncoder' in paragraph.text:
                issues.append("❌ Found LabelEncoder mention (should be OneHotEncoder)")
                break
        
        # Check for wrong accuracy values
        for paragraph in doc.paragraphs:
            if '74.2%' in paragraph.text or '74.0%' in paragraph.text:
                issues.append("⚠️ Found old accuracy values (74.2%, 74.0%)")
                break
        
        # Check for wrong feature count
        for paragraph in doc.paragraphs:
            if '19 features' in paragraph.text:
                issues.append("⚠️ Found old feature count (19, should be 21)")
                break
        
        if issues:
            print("\n=== ISSUES FOUND ===")
            for issue in issues:
                print(issue)
        else:
            print("\n✅ No major issues found!")
        
        # Overall assessment
        print("\n=== OVERALL ASSESSMENT ===")
        if accuracy_found and model_found and features_found and preprocessing_found and not issues:
            print("✅ Report appears to be well synced with project results")
        elif len(issues) == 0:
            print("⚠️ Report is mostly synced but may need minor updates")
        else:
            print("❌ Report needs updates to match project results")
        
    except Exception as e:
        print(f"Error analyzing report: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    analyze_report_content()
